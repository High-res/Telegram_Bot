from docx import Document
from docx.shared import Pt
from .getAvans import getAvanses
from .getWorker import getWorkerss
from .sendDocx import send_docx




async def addToDocx() :
    document = Document('zayavlenie_avans.DOCX')

    paragraphs = document.paragraphs
    style = document.styles['Normal']
    font = style.font
    to_avans = []
    last_req = getAvanses()[-1]
    all_users = getWorkerss()
    for i in range(0, len(getWorkerss())):
        if int(all_users[i]['tg_id']) == int(last_req['user_name']) :
            to_avans.append(all_users[i])

    print(last_req)

    year_day, separator, day_time = last_req['date'].partition(' ')


    paragraphs[0].add_run('Директору ТОО «UCJ.KZ»  ').bold = True
    paragraphs[1].add_run('Бибитбеку А.Б.  ').bold = True
    paragraphs[2].add_run(f'От ' +  to_avans[0]['name']).bold = True
    paragraphs[14].add_run(f'              Прошу вас выдать аванс в размере ' + last_req['answer'] + ' тг в счет заработной платы за ' + year_day.split('-')[1] + ' месяц 2021г')
    paragraphs[24].add_run(f'Дата: ' + year_day +'                                                     Подпись _________________')
    font.name = 'Times New Roman'
    font.size = Pt(11)
    doc_name = f''+to_avans[0]['name']+' ' + year_day +'.docx'
    document.save(f''+to_avans[0]['name']+' ' + year_day +'.docx')
    await send_docx(doc_name)
